#!/usr/bin/env python3
"""Build the V4 course outline DOCX from the approved HTML content."""

from __future__ import annotations

import sys
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


class Node:
    def __init__(self, tag: str = "", attrs=None, parent=None, text: str = ""):
        self.tag = tag
        self.attrs = dict(attrs or [])
        self.parent = parent
        self.text = text
        self.children: list[Node] = []


class BodyParser(HTMLParser):
    VOID_TAGS = {"meta", "link", "img", "br", "hr", "input"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = Node("root")
        self.stack = [self.root]
        self.in_body = False

    def handle_starttag(self, tag, attrs):
        if tag == "body":
            self.in_body = True
            return
        if not self.in_body:
            return
        node = Node(tag, attrs, self.stack[-1])
        self.stack[-1].children.append(node)
        if tag not in self.VOID_TAGS:
            self.stack.append(node)

    def handle_endtag(self, tag):
        if tag == "body":
            self.in_body = False
            return
        if not self.in_body:
            return
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                break

    def handle_data(self, data):
        if self.in_body and data:
            self.stack[-1].children.append(Node("#text", parent=self.stack[-1], text=data))


BLUE = "1F4E79"
MID_BLUE = "5B9BD5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "F3F7FB"
LIGHT_GREEN = "E2F0D9"
LIGHT_YELLOW = "FFF2CC"
LIGHT_PURPLE = "E4DFEC"
TEXT = RGBColor(34, 34, 34)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_font(run, size=10.5, bold=False, color=None, name="PingFang SC"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color or TEXT


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def plain_text(node: Node) -> str:
    if node.tag == "#text":
        return node.text
    return "".join(plain_text(child) for child in node.children)


def normalized_text(node: Node) -> str:
    return " ".join(plain_text(node).split())


def add_inline(paragraph, node: Node, inherited_bold=False):
    if node.tag == "#text":
        if node.text:
            run = paragraph.add_run(node.text)
            set_font(run, bold=inherited_bold)
        return
    bold = inherited_bold or node.tag in {"b", "strong"}
    for child in node.children:
        if child.tag == "#text":
            text = child.text
            if not text:
                continue
            run = paragraph.add_run(text)
            fill = None
            if node.tag == "span":
                klass = node.attrs.get("class", "")
                fill = {"tag": LIGHT_GREEN, "warn": LIGHT_YELLOW, "tool": LIGHT_PURPLE}.get(klass)
            if node.tag == "code":
                run.font.name = "Menlo"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
                run.font.size = Pt(9)
            else:
                set_font(run, bold=bold)
            if fill:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), fill)
                run._r.get_or_add_rPr().append(shd)
        else:
            add_inline(paragraph, child, bold)


def style_paragraph(paragraph, after=5, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_body_paragraph(doc, node, style=None, left_indent=None):
    p = doc.add_paragraph(style=style)
    add_inline(p, node)
    style_paragraph(p)
    if left_indent:
        p.paragraph_format.left_indent = Cm(left_indent)
    return p


def direct_children(node: Node, tags: set[str]) -> list[Node]:
    return [child for child in node.children if child.tag in tags]


def add_table(doc, node: Node):
    row_nodes = []
    for child in node.children:
        if child.tag == "tr":
            row_nodes.append(child)
        elif child.tag in {"thead", "tbody", "tfoot"}:
            row_nodes.extend(direct_children(child, {"tr"}))
    if not row_nodes:
        return
    parsed_rows = []
    max_cols = 0
    for row_node in row_nodes:
        cells = direct_children(row_node, {"th", "td"})
        parsed_rows.append(cells)
        max_cols = max(max_cols, len(cells))
    table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ridx, cells in enumerate(parsed_rows):
        row = table.rows[ridx]
        set_cant_split(row)
        if ridx == 0:
            set_repeat_table_header(row)
        for cidx in range(max_cols):
            cell = row.cells[cidx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if cidx < len(cells):
                add_inline(p, cells[cidx], inherited_bold=(ridx == 0 or cells[cidx].tag == "th"))
            for run in p.runs:
                set_font(run, size=9.0, bold=(ridx == 0 or run.bold))
            if ridx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_quote(doc, node: Node):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=130, start=180, bottom=130, end=180)
    p = cell.paragraphs[0]
    add_inline(p, node)
    style_paragraph(p, after=0, line=1.35)
    for run in p.runs:
        set_font(run, size=10.5, color=RGBColor(47, 85, 151))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_list(doc, node: Node, ordered=False):
    idx = 0
    for child in node.children:
        if child.tag != "li":
            continue
        idx += 1
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        add_inline(p, child)
        style_paragraph(p, after=3, line=1.28)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("第一阶段 · AI产品经理系统学习班 · V4.0")
    set_font(run, size=8.5, color=RGBColor(127, 127, 127))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("执行与售卖版  ·  ")
    set_font(run, size=8.5, color=RGBColor(127, 127, 127))
    add_page_field(footer)


def add_cover(doc: Document, title: str, subtitle: str):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_font(run, size=25, bold=True, color=RGBColor(23, 54, 93))
    p.paragraph_format.space_after = Pt(22)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("AI PRODUCT MANAGER · FOUNDATION")
    set_font(run, size=9.5, color=RGBColor(91, 155, 213), name="Arial")
    line.paragraph_format.space_after = Pt(32)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, part in enumerate(subtitle.split("｜")):
        if index:
            sub.add_run("\n")
        run = sub.add_run(part.strip())
        set_font(run, size=13 if index == 0 else 11, bold=(index == 0), color=RGBColor(68, 84, 106))
    sub.paragraph_format.line_spacing = 1.5

    for _ in range(5):
        doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("课程内部执行基准｜核心知识可追溯｜案例均为公开信息或教学虚构")
    set_font(run, size=9, color=RGBColor(127, 127, 127))
    doc.add_page_break()


def add_heading(doc, node: Node, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(15 if level == 2 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 2 else 5)
    run = p.add_run(normalized_text(node))
    if level == 2:
        set_font(run, size=16, bold=True, color=RGBColor(31, 78, 121))
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), MID_BLUE)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        set_font(run, size=13.5, bold=True, color=RGBColor(47, 85, 151))
        p.paragraph_format.keep_with_next = True
    else:
        set_font(run, size=11.5, bold=True, color=RGBColor(68, 84, 106))


def build(html_path: Path, output_path: Path):
    parser = BodyParser()
    parser.feed(html_path.read_text(encoding="utf-8"))
    body_children = parser.root.children

    title_node = next(node for node in body_children if node.tag == "h1")
    title_index = body_children.index(title_node)
    subtitle_node = next(
        node for node in body_children[title_index + 1 :]
        if node.tag == "p" and normalized_text(node)
    )

    doc = Document()
    configure_document(doc)
    add_cover(doc, normalized_text(title_node), normalized_text(subtitle_node))

    skipped = {id(title_node), id(subtitle_node)}
    for node in body_children:
        if id(node) in skipped or node.tag == "#text":
            continue
        if node.tag in {"h2", "h3", "h4"}:
            add_heading(doc, node, int(node.tag[1]))
        elif node.tag == "p":
            p = add_body_paragraph(doc, node)
            klass = node.attrs.get("class", "")
            if klass in {"small", "note"}:
                for run in p.runs:
                    set_font(run, size=9.2, color=RGBColor(85, 85, 85))
            if klass == "note":
                p.paragraph_format.left_indent = Cm(0.25)
                p.paragraph_format.right_indent = Cm(0.25)
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_YELLOW)
                pPr.append(shd)
        elif node.tag == "blockquote":
            add_quote(doc, node)
        elif node.tag == "table":
            add_table(doc, node)
        elif node.tag == "ul":
            add_list(doc, node, ordered=False)
        elif node.tag == "ol":
            add_list(doc, node, ordered=True)

    doc.core_properties.title = "第一阶段：AI产品经理系统学习班｜V4.0执行与售卖版"
    doc.core_properties.subject = "第0课＋12次正式课的课程总大纲"
    doc.core_properties.author = "课程设计团队"
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_v4_outline_docx.py INPUT.html OUTPUT.docx")
    build(Path(sys.argv[1]).resolve(), Path(sys.argv[2]).resolve())
