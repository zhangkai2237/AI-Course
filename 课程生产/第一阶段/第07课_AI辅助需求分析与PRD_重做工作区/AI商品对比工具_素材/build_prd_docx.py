from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "商品对比工具_PRD.md"
OUTPUT = ROOT / "商品对比工具_PRD_V1.0.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
MUTED = "667085"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
BORDER = "D0D5DD"
WHITE = "FFFFFF"
FONT_LATIN = "Calibri"
FONT_CJK = "PingFang SC"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def add_bottom_border(paragraph, color=BLUE, size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_inline_markdown(paragraph, text, default_size=10.5, default_color=None):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if bold else part
        run = paragraph.add_run(content)
        set_run_font(run, size=default_size, bold=bold, color=default_color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("202124")
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("产品需求文档 · PRD")
    set_run_font(r, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("商品对比工具")
    set_run_font(r, size=26, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("帮助用户更快看懂差异、收敛候选并完成购买决策")
    set_run_font(r, size=13, color=MUTED)

    metadata = [
        ("文档版本", "V1.0"),
        ("文档状态", "评审稿"),
        ("产品阶段", "MVP → 增强版"),
        ("适用端", "Web / H5 / App"),
        ("更新时间", "2026-08-09"),
        ("责任人", "产品 / 设计 / 研发 / 测试 / 数据：待定"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        a = p.add_run(f"{label}：")
        set_run_font(a, size=10.5, bold=True, color=INK)
        b = p.add_run(value)
        set_run_font(b, size=10.5, color="344054")

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(12)
    rule.paragraph_format.space_after = Pt(16)
    add_bottom_border(rule, BLUE, "14")

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    set_table_borders(callout, color="B9D6EA", size="8")
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("核心主张")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    add_inline_markdown(
        p2,
        "通过统一属性口径、突出关键差异并连接实时购买信息，把多页面的信息搜集任务变成一次可解释、可行动的横向决策。",
        10.5,
        "202124",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("阅读指南")
    set_run_font(r, size=11, bold=True, color=INK)
    p = doc.add_paragraph(style="List Bullet")
    add_inline_markdown(p, "产品评审：重点阅读第 1–6、14、19 节。")
    p = doc.add_paragraph(style="List Bullet")
    add_inline_markdown(p, "研发与测试：重点阅读第 6–10、14–16 节。")
    p = doc.add_paragraph(style="List Bullet")
    add_inline_markdown(p, "数据与运营：重点阅读第 2、7、11–13、17 节。")

def set_running_furniture(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    left = p.add_run("商品对比工具 PRD")
    set_run_font(left, size=9, bold=True, color=MUTED)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = fp.add_run("内部评审稿 · V1.0")
    set_run_font(footer_run, size=9, color=MUTED)


def add_markdown_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        for cidx, value in enumerate(row_data):
            cell = row.cells[cidx]
            if ridx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_inline_markdown(p, value.strip(), default_size=9.2, default_color=INK if ridx == 0 else "202124")
            for run in p.runs:
                if ridx == 0:
                    run.bold = True
    if col_count == 2:
        widths = [2700, 6660]
    elif col_count == 3:
        widths = [2100, 3630, 3630]
    elif col_count == 4:
        widths = [1700, 2660, 2500, 2500]
    else:
        base = 9360 // col_count
        widths = [base] * col_count
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    seen_title = False
    first_body_heading = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# ") and not seen_title:
            seen_title = True
            i += 1
            continue
        if stripped.startswith(">") and i < 5:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue
        if stripped.startswith("|"):
            raw_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
                raw_rows.append(parts)
                i += 1
            rows = [row for row in raw_rows if not all(re.fullmatch(r":?-{3,}:?", c or "") for c in row)]
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            if level == 1 and first_body_heading:
                p.paragraph_format.page_break_before = True
                first_body_heading = False
            add_inline_markdown(p, heading.group(2), default_size={1: 16, 2: 13, 3: 11.5}[level], default_color=BLUE if level < 3 else DARK_BLUE)
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, bullet.group(1))
            i += 1
            continue
        number = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number:
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, number.group(1))
            i += 1
            continue
        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)
        i += 1


def add_document_properties(doc):
    props = doc.core_properties
    props.title = "商品对比工具 PRD"
    props.subject = "电商平台商品对比工具完整产品需求文档"
    props.author = "产品团队"
    props.keywords = "电商, 商品对比, PRD, 购买决策"
    props.comments = "基于用户给出的单句业务目标从零生成。"


def main():
    doc = Document()
    configure_section(doc.sections[0])
    configure_styles(doc)
    set_running_furniture(doc)
    add_document_properties(doc)
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
